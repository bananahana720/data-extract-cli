"""DOCX extractor adapter."""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Tuple

from data_extract.extract.adapter import ExtractorAdapter


class DocxExtractorAdapter(ExtractorAdapter):
    """Extractor for Microsoft Word documents."""

    def __init__(self) -> None:
        super().__init__(format_name="docx")

    def extract(self, file_path: Path) -> Tuple[str, Dict[str, Any], Dict[str, float]]:
        try:
            from docx import Document as WordDocument
        except ImportError as exc:
            raise RuntimeError("python-docx is required for DOCX extraction") from exc

        doc = WordDocument(str(file_path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]

        table_lines = []
        for table in doc.tables:
            for row in table.rows:
                table_lines.append(" | ".join(cell.text.strip() for cell in row.cells))

        text_parts = paragraphs + table_lines
        text = "\n\n".join(part for part in text_parts if part)

        structure = {
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "char_count": len(text),
        }
        quality = {
            "extraction_confidence": 1.0,
        }
        return text, structure, quality
