"""Generate synthetic test corpus for UAT testing.

Creates deterministic, PII-free test documents for validating
the data extraction pipeline and CLI user journeys.

Creates:
- 5 PDF files using reportlab (simple text PDFs)
- 3 DOCX files using python-docx
- 2 XLSX files using openpyxl

All content is synthetic, deterministic, and covers varied topics
relevant to enterprise document processing (audit, compliance, financial).

Story 5-0: UAT Testing Framework
Usage:
    python tests/uat/fixtures/generate_corpus.py
"""

from __future__ import annotations

import sys
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    pass


def check_dependencies() -> dict[str, bool]:
    """Check availability of optional dependencies.

    Returns:
        dict: Availability status for each dependency
    """
    dependencies = {}

    try:
        import reportlab  # type: ignore[import-untyped]  # noqa: F401

        dependencies["reportlab"] = True
    except ImportError:
        dependencies["reportlab"] = False

    try:
        import docx  # type: ignore[import-untyped]  # noqa: F401

        dependencies["docx"] = True
    except ImportError:
        dependencies["docx"] = False

    try:
        import openpyxl  # type: ignore[import-untyped]  # noqa: F401

        dependencies["openpyxl"] = True
    except ImportError:
        dependencies["openpyxl"] = False

    return dependencies


def create_pdf_files(output_dir: Path) -> list[Path]:
    """Create 5 sample PDF files using reportlab.

    Args:
        output_dir: Directory to save PDF files

    Returns:
        list[Path]: Paths to created PDF files
    """
    try:
        from reportlab.lib.pagesizes import letter  # type: ignore[import-untyped]
        from reportlab.lib.styles import getSampleStyleSheet  # type: ignore[import-untyped]
        from reportlab.lib.units import inch  # type: ignore[import-untyped]
        from reportlab.platypus import (  # type: ignore[import-untyped]
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
    except ImportError:
        print("⚠ reportlab not available - skipping PDF generation", file=sys.stderr)
        print("  Install with: pip install reportlab", file=sys.stderr)
        return []

    created_files = []
    styles = getSampleStyleSheet()

    # PDF 1: Audit Report - Financial Controls
    pdf1_path = output_dir / "audit-financial-controls-q3-2024.pdf"
    pdf1 = SimpleDocTemplate(str(pdf1_path), pagesize=letter)
    pdf1_content = [
        Paragraph("Internal Audit Report: Financial Controls Q3 2024", styles["Title"]),
        Spacer(1, 0.2 * inch),
        Paragraph("Executive Summary", styles["Heading1"]),
        Paragraph(
            "This audit assessed the effectiveness of financial controls across "
            "three divisions: operations, sales, and corporate finance. Our testing "
            "covered 847 transactions totaling $12.4 million in expenditures.",
            styles["BodyText"],
        ),
        Spacer(1, 0.1 * inch),
        Paragraph("Key Findings", styles["Heading2"]),
        Paragraph(
            "1. Segregation of duties: Adequate controls in place for approval workflows. "
            "2. Invoice processing: Identified 23 invoices processed without proper documentation. "
            "3. Expense reimbursement: Policy compliance rate of 94%, up from 87% in Q2. "
            "4. Reconciliation procedures: Monthly bank reconciliations completed timely in 11/12 accounts.",
            styles["BodyText"],
        ),
        Spacer(1, 0.1 * inch),
        Paragraph("Recommendations", styles["Heading2"]),
        Paragraph(
            "Implement automated controls for invoice documentation validation. "
            "Enhance training for expense policy compliance. "
            "Review reconciliation process for the one outlier account.",
            styles["BodyText"],
        ),
    ]
    pdf1.build(pdf1_content)
    created_files.append(pdf1_path)

    # PDF 2: Compliance Review - Data Privacy
    pdf2_path = output_dir / "compliance-data-privacy-review.pdf"
    pdf2 = SimpleDocTemplate(str(pdf2_path), pagesize=letter)
    pdf2_content = [
        Paragraph("Compliance Review: Data Privacy Controls", styles["Title"]),
        Spacer(1, 0.2 * inch),
        Paragraph("Scope and Methodology", styles["Heading1"]),
        Paragraph(
            "Reviewed data handling practices for customer information across "
            "all departments. Assessed compliance with GDPR, CCPA, and internal "
            "data governance policies. Testing included 156 data processing activities.",
            styles["BodyText"],
        ),
        Spacer(1, 0.1 * inch),
        Paragraph("Compliance Findings", styles["Heading2"]),
        Paragraph(
            "Data minimization: 89% of systems collect only necessary data. "
            "Access controls: Role-based access properly configured in 94% of systems. "
            "Retention policies: Automated deletion implemented for 67% of data categories. "
            "Consent management: Valid consent documented for 98% of marketing contacts.",
            styles["BodyText"],
        ),
        Spacer(1, 0.1 * inch),
        Paragraph("Action Items", styles["Heading2"]),
        Paragraph(
            "Extend automated deletion to remaining data categories by Q1 2025. "
            "Review systems with unnecessary data collection. "
            "Update consent forms for 2% of contacts missing documentation.",
            styles["BodyText"],
        ),
    ]
    pdf2.build(pdf2_content)
    created_files.append(pdf2_path)

    # PDF 3: Risk Assessment - Operational Risk
    pdf3_path = output_dir / "risk-assessment-operational.pdf"
    pdf3 = SimpleDocTemplate(str(pdf3_path), pagesize=letter)
    pdf3_content = [
        Paragraph("Operational Risk Assessment 2024", styles["Title"]),
        Spacer(1, 0.2 * inch),
        Paragraph("Risk Landscape Overview", styles["Heading1"]),
        Paragraph(
            "Assessed 42 operational risk scenarios across supply chain, technology, "
            "human resources, and facilities management. Risk ratings based on "
            "likelihood and impact scoring methodology.",
            styles["BodyText"],
        ),
        Spacer(1, 0.1 * inch),
        Paragraph("High-Priority Risks", styles["Heading2"]),
        Paragraph(
            "1. Technology infrastructure: Aging systems pose continuity risk (High). "
            "2. Vendor concentration: 60% of critical services from 3 vendors (Medium-High). "
            "3. Skills gap: Key technical roles with single points of failure (High). "
            "4. Facility security: Physical access controls need modernization (Medium).",
            styles["BodyText"],
        ),
        Spacer(1, 0.1 * inch),
        Paragraph("Mitigation Strategy", styles["Heading2"]),
        Paragraph(
            "Technology modernization roadmap approved for 2025-2026. "
            "Vendor diversification initiative launched. "
            "Succession planning and cross-training programs expanded. "
            "Facility security upgrade budgeted for Q2 2025.",
            styles["BodyText"],
        ),
    ]
    pdf3.build(pdf3_content)
    created_files.append(pdf3_path)

    # PDF 4: Financial Analysis - Budget Variance
    pdf4_path = output_dir / "financial-analysis-budget-variance.pdf"
    pdf4 = SimpleDocTemplate(str(pdf4_path), pagesize=letter)
    pdf4_content = [
        Paragraph("Budget Variance Analysis: Year-End 2024", styles["Title"]),
        Spacer(1, 0.2 * inch),
        Paragraph("Performance Summary", styles["Heading1"]),
        Paragraph(
            "Total operating budget: $8.9M. Actual expenditures: $8.7M. "
            "Overall variance: 2.2% under budget. Detailed analysis by department "
            "reveals mixed performance across cost centers.",
            styles["BodyText"],
        ),
        Spacer(1, 0.1 * inch),
        Paragraph("Departmental Variance", styles["Heading2"]),
        Paragraph(
            "Operations: 5.3% under budget due to efficiency gains and delayed hiring. "
            "Marketing: 3.8% over budget from additional campaign spending. "
            "Technology: 1.2% under budget with deferred infrastructure projects. "
            "Administration: 0.5% over budget from professional services fees.",
            styles["BodyText"],
        ),
        Spacer(1, 0.1 * inch),
        Paragraph("Outlook", styles["Heading2"]),
        Paragraph(
            "2025 budget planning incorporates lessons from variance analysis. "
            "Marketing budget increased by 5% for strategic initiatives. "
            "Technology infrastructure projects prioritized in Q1-Q2. "
            "Hiring plan adjusted based on operational efficiency data.",
            styles["BodyText"],
        ),
    ]
    pdf4.build(pdf4_content)
    created_files.append(pdf4_path)

    # PDF 5: Meeting Minutes - Board Meeting
    pdf5_path = output_dir / "board-meeting-minutes-2024-11.pdf"
    pdf5 = SimpleDocTemplate(str(pdf5_path), pagesize=letter)
    pdf5_content = [
        Paragraph("Board of Directors Meeting Minutes", styles["Title"]),
        Paragraph("November 15, 2024", styles["Heading2"]),
        Spacer(1, 0.2 * inch),
        Paragraph("Attendees", styles["Heading1"]),
        Paragraph(
            "Present: Directors Smith, Johnson, Chen, Williams, Davis. "
            "Absent: Director Martinez (excused). "
            "Management: CEO Roberts, CFO Taylor, COO Anderson.",
            styles["BodyText"],
        ),
        Spacer(1, 0.1 * inch),
        Paragraph("Financial Report", styles["Heading2"]),
        Paragraph(
            "CFO Taylor presented Q3 financial results. Revenue of $23.4M, "
            "up 8% YoY. Operating margin of 14.2%, consistent with plan. "
            "Cash position strong at $12.8M. Board approved annual audit engagement.",
            styles["BodyText"],
        ),
        Spacer(1, 0.1 * inch),
        Paragraph("Strategic Initiatives", styles["Heading2"]),
        Paragraph(
            "CEO Roberts reviewed progress on digital transformation initiative. "
            "Customer portal launch scheduled for Q1 2025. "
            "Process automation project on track, 35% complete. "
            "Board approved capital allocation for technology investments.",
            styles["BodyText"],
        ),
        Spacer(1, 0.1 * inch),
        Paragraph("Next Meeting", styles["Heading2"]),
        Paragraph("Scheduled for February 14, 2025 at 2:00 PM.", styles["BodyText"]),
    ]
    pdf5.build(pdf5_content)
    created_files.append(pdf5_path)

    return created_files


def create_docx_files(output_dir: Path) -> list[Path]:
    """Create 3 sample DOCX files using python-docx.

    Args:
        output_dir: Directory to save DOCX files

    Returns:
        list[Path]: Paths to created DOCX files
    """
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError:
        print("⚠ python-docx not available - skipping DOCX generation", file=sys.stderr)
        print("  Install with: pip install python-docx", file=sys.stderr)
        return []

    created_files = []

    # DOCX 1: Policy Document - Remote Work
    docx1_path = output_dir / "policy-remote-work-guidelines.docx"
    doc1 = Document()
    doc1.add_heading("Remote Work Policy Guidelines", level=1)
    doc1.add_heading("Purpose", level=2)
    doc1.add_paragraph(
        "This policy establishes guidelines for remote work arrangements to "
        "ensure productivity, security, and work-life balance for all employees. "
        "Applies to full-time, part-time, and contract workers."
    )
    doc1.add_heading("Eligibility", level=2)
    doc1.add_paragraph(
        "Employees may request remote work if: (1) Role responsibilities can be "
        "performed remotely, (2) Performance meets or exceeds expectations, "
        "(3) Home workspace meets ergonomic and security requirements."
    )
    doc1.add_heading("Requirements", level=2)
    doc1.add_paragraph(
        "Equipment: Company-issued laptop with encrypted storage. "
        "Connectivity: Minimum 25 Mbps internet connection. "
        "Security: Use of VPN for all company system access. "
        "Availability: Core hours 10 AM - 3 PM local time for team collaboration."
    )
    doc1.add_heading("Approval Process", level=2)
    doc1.add_paragraph(
        "Submit remote work request to manager with completed workspace assessment form. "
        "Manager approval required within 5 business days. "
        "IT department conducts security review before access granted."
    )
    doc1.save(str(docx1_path))
    created_files.append(docx1_path)

    # DOCX 2: Procedure Document - Expense Reimbursement
    docx2_path = output_dir / "procedure-expense-reimbursement.docx"
    doc2 = Document()
    doc2.add_heading("Expense Reimbursement Procedure", level=1)
    doc2.add_heading("Overview", level=2)
    doc2.add_paragraph(
        "Employees may seek reimbursement for approved business expenses incurred "
        "in the course of their duties. This procedure outlines submission requirements "
        "and processing timelines."
    )
    doc2.add_heading("Eligible Expenses", level=2)
    doc2.add_paragraph(
        "Travel: Airfare, lodging, ground transportation (excluding first-class upgrades). "
        "Meals: Up to $75 per day while traveling on business. "
        "Supplies: Office supplies and materials for business use. "
        "Client Entertainment: Reasonable meals and activities with prior approval."
    )
    doc2.add_heading("Submission Requirements", level=2)
    doc2.add_paragraph(
        "1. Complete expense report form within 30 days of expense date. "
        "2. Attach original itemized receipts (required for expenses over $25). "
        "3. Include business purpose and attendees for meals/entertainment. "
        "4. Obtain manager approval signature before submitting to Finance."
    )
    doc2.add_heading("Processing Timeline", level=2)
    doc2.add_paragraph(
        "Finance reviews submissions within 5 business days. "
        "Approved reimbursements processed with next bi-weekly payroll. "
        "Incomplete submissions returned to employee with explanation. "
        "Questions directed to finance@example.com."
    )
    doc2.save(str(docx2_path))
    created_files.append(docx2_path)

    # DOCX 3: Project Summary - System Modernization
    docx3_path = output_dir / "project-summary-system-modernization.docx"
    doc3 = Document()
    doc3.add_heading("Project Summary: System Modernization Initiative", level=1)
    doc3.add_heading("Project Charter", level=2)
    doc3.add_paragraph(
        "Project Name: Core Systems Modernization. "
        "Sponsor: CIO Rebecca Chen. "
        "Duration: 18 months (January 2025 - June 2026). "
        "Budget: $2.4M capital, $800K operational."
    )
    doc3.add_heading("Objectives", level=2)
    doc3.add_paragraph(
        "1. Replace legacy ERP system with cloud-based solution. "
        "2. Implement modern CRM with customer portal. "
        "3. Migrate on-premise infrastructure to hybrid cloud. "
        "4. Establish automated backup and disaster recovery."
    )
    doc3.add_heading("Key Milestones", level=2)
    doc3.add_paragraph(
        "Phase 1 (Q1 2025): Requirements gathering and vendor selection. "
        "Phase 2 (Q2-Q3 2025): ERP implementation and data migration. "
        "Phase 3 (Q4 2025): CRM deployment and user training. "
        "Phase 4 (Q1-Q2 2026): Cloud migration and final testing."
    )
    doc3.add_heading("Success Metrics", level=2)
    doc3.add_paragraph(
        "System uptime: 99.9% availability. "
        "User adoption: 95% active usage within 90 days of go-live. "
        "Performance: 50% reduction in report generation time. "
        "Cost: Total cost of ownership reduction of 30% over 3 years."
    )
    doc3.save(str(docx3_path))
    created_files.append(docx3_path)

    return created_files


def create_xlsx_files(output_dir: Path) -> list[Path]:
    """Create 2 sample XLSX files using openpyxl.

    Args:
        output_dir: Directory to save XLSX files

    Returns:
        list[Path]: Paths to created XLSX files
    """
    try:
        from openpyxl import Workbook  # type: ignore[import-untyped]
        from openpyxl.styles import Font  # type: ignore[import-untyped]
    except ImportError:
        print("⚠ openpyxl not available - skipping XLSX generation", file=sys.stderr)
        print("  Install with: pip install openpyxl", file=sys.stderr)
        return []

    created_files = []

    # XLSX 1: Financial Summary - Quarterly Results
    xlsx1_path = output_dir / "financial-summary-q3-2024.xlsx"
    wb1 = Workbook()
    ws1 = wb1.active
    assert ws1 is not None  # Active sheet always exists for new workbook
    ws1.title = "Q3 Summary"

    # Headers
    ws1["A1"] = "Account"
    ws1["B1"] = "Budget"
    ws1["C1"] = "Actual"
    ws1["D1"] = "Variance"
    ws1["E1"] = "Variance %"
    for cell_ref in ["A1", "B1", "C1", "D1", "E1"]:
        ws1[cell_ref].font = Font(bold=True)

    # Data rows
    data: list[list[str | int]] = [
        ["Revenue", 6500000, 6890000, 390000, "6.0%"],
        ["Cost of Sales", 2800000, 2650000, -150000, "-5.4%"],
        ["Gross Profit", 3700000, 4240000, 540000, "14.6%"],
        ["Operating Expenses", 2900000, 2876000, -24000, "-0.8%"],
        ["Operating Income", 800000, 1364000, 564000, "70.5%"],
        ["Interest Expense", 45000, 42000, -3000, "-6.7%"],
        ["Net Income", 755000, 1322000, 567000, "75.1%"],
    ]
    for idx, row in enumerate(data, start=2):
        for col_idx, value in enumerate(row, start=1):
            ws1.cell(row=idx, column=col_idx, value=value)  # type: ignore[arg-type]

    # Notes sheet
    ws2 = wb1.create_sheet("Notes")
    ws2["A1"] = "Notes and Assumptions"
    ws2["A1"].font = Font(bold=True)
    ws2["A3"] = "1. Revenue exceeds budget due to strong Q3 sales performance"
    ws2["A4"] = "2. Cost of sales lower from improved vendor pricing"
    ws2["A5"] = "3. Operating expenses well-controlled, minimal variance"
    ws2["A6"] = "4. Operating income significantly above plan"

    wb1.save(str(xlsx1_path))
    created_files.append(xlsx1_path)

    # XLSX 2: Audit Findings - Control Testing
    xlsx2_path = output_dir / "audit-findings-control-testing.xlsx"
    wb2 = Workbook()
    ws_findings = wb2.active
    assert ws_findings is not None  # Active sheet always exists for new workbook
    ws_findings.title = "Findings"

    # Headers
    headers = ["Finding ID", "Control Area", "Risk Level", "Status", "Owner", "Due Date"]
    for col_idx, header in enumerate(headers, start=1):
        header_cell = ws_findings.cell(row=1, column=col_idx, value=header)
        header_cell.font = Font(bold=True)

    # Findings data
    findings: list[list[str]] = [
        ["F-001", "Invoice Approval", "Medium", "Open", "A. Johnson", "2024-12-15"],
        ["F-002", "Bank Reconciliation", "Low", "Closed", "B. Smith", "2024-11-30"],
        ["F-003", "Access Controls", "High", "In Progress", "C. Davis", "2024-12-31"],
        ["F-004", "Documentation", "Low", "Open", "D. Wilson", "2025-01-15"],
        ["F-005", "Segregation of Duties", "Medium", "In Progress", "E. Martinez", "2024-12-20"],
        ["F-006", "Change Management", "Medium", "Open", "F. Garcia", "2025-01-10"],
        ["F-007", "Password Policy", "Low", "Closed", "G. Lee", "2024-11-15"],
        ["F-008", "Data Backup", "High", "In Progress", "H. Chen", "2024-12-10"],
    ]
    for row_idx, finding in enumerate(findings, start=2):
        for col_idx, value in enumerate(finding, start=1):
            ws_findings.cell(row=row_idx, column=col_idx, value=value)

    # Summary sheet
    ws2 = wb2.create_sheet("Summary")
    ws2["A1"] = "Findings Summary"
    ws2["A1"].font = Font(bold=True)
    ws2["A3"] = "Total Findings: 8"
    ws2["A4"] = "High Risk: 2"
    ws2["A5"] = "Medium Risk: 4"
    ws2["A6"] = "Low Risk: 2"
    ws2["A8"] = "Status Breakdown:"
    ws2["A9"] = "Open: 3"
    ws2["A10"] = "In Progress: 3"
    ws2["A11"] = "Closed: 2"

    wb2.save(str(xlsx2_path))
    created_files.append(xlsx2_path)

    return created_files


def main() -> int:
    """Generate all test corpus files.

    Returns:
        int: Exit code (0 for success, 1 for partial success, 2 for failure)
    """
    # Determine output directory
    script_dir = Path(__file__).parent
    output_dir = script_dir / "sample_corpus"
    output_dir.mkdir(parents=True, exist_ok=True)

    print(f"📁 Generating test corpus in: {output_dir}")
    print()

    # Check dependencies
    deps = check_dependencies()
    print("📦 Dependency check:")
    for dep, available in deps.items():
        status = "✓" if available else "✗"
        print(f"  {status} {dep}")
    print()

    # Generate files
    all_created = []
    error_count = 0

    # PDFs
    if deps["reportlab"]:
        try:
            pdf_files = create_pdf_files(output_dir)
            all_created.extend(pdf_files)
            print(f"✓ Created {len(pdf_files)} PDF files")
        except Exception as e:
            print(f"✗ PDF generation failed: {e}", file=sys.stderr)
            error_count += 1
    else:
        print("⊘ Skipped PDF generation (reportlab not available)")
        error_count += 1

    # DOCX
    if deps["docx"]:
        try:
            docx_files = create_docx_files(output_dir)
            all_created.extend(docx_files)
            print(f"✓ Created {len(docx_files)} DOCX files")
        except Exception as e:
            print(f"✗ DOCX generation failed: {e}", file=sys.stderr)
            error_count += 1
    else:
        print("⊘ Skipped DOCX generation (python-docx not available)")
        error_count += 1

    # XLSX
    if deps["openpyxl"]:
        try:
            xlsx_files = create_xlsx_files(output_dir)
            all_created.extend(xlsx_files)
            print(f"✓ Created {len(xlsx_files)} XLSX files")
        except Exception as e:
            print(f"✗ XLSX generation failed: {e}", file=sys.stderr)
            error_count += 1
    else:
        print("⊘ Skipped XLSX generation (openpyxl not available)")
        error_count += 1

    print()
    print(f"📊 Total files created: {len(all_created)}")

    if all_created:
        print()
        print("Created files:")
        for file in all_created:
            print(f"  • {file.name}")

    # Generate summary file
    summary_path = output_dir / "CORPUS_INFO.txt"
    with open(summary_path, "w") as f:
        f.write("UAT Test Corpus\n")
        f.write("=" * 50 + "\n\n")
        f.write(f"Generated: {datetime.now().isoformat()}\n")
        f.write(f"Total files: {len(all_created)}\n\n")
        f.write("File Details:\n")
        f.write("-" * 50 + "\n")
        for file in sorted(all_created):
            size_kb = file.stat().st_size / 1024
            f.write(f"{file.name:50} {size_kb:8.1f} KB\n")
        f.write("\n")
        f.write("Content Notes:\n")
        f.write("-" * 50 + "\n")
        f.write("All content is synthetic and PII-free.\n")
        f.write("Topics: audit, compliance, financial, risk, policy, procedures.\n")
        f.write("Designed for deterministic UAT testing.\n")

    print(f"\n📄 Corpus info saved to: {summary_path.name}")

    # Determine exit code
    if error_count == 0:
        print("\n✓ Corpus generation complete - all formats successful")
        return 0
    elif len(all_created) > 0:
        print(f"\n⚠ Corpus generation partial - {error_count} format(s) failed", file=sys.stderr)
        return 1
    else:
        print("\n✗ Corpus generation failed - no files created", file=sys.stderr)
        print("  Install dependencies: pip install reportlab python-docx openpyxl", file=sys.stderr)
        return 2


if __name__ == "__main__":
    sys.exit(main())
